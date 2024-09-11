from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.contrib import messages
import pandas as pd
from .models import QuestionBank,InputSuggestion,InputSuggestionImage, InputSuggestionDocument, ExamName, Subject, Area, PartName, ChapterName
from django.db.models import Max
from .forms import UploadFileForm
import os
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from django.http import FileResponse, HttpResponse
from PIL import Image as PILImage
from django.utils.text import slugify
from io import BytesIO
from django.conf import settings
from django.http import HttpResponseServerError
from django.http import JsonResponse
from .models import Subject, Area, PartName, TopicName
# ************************* Generate Test Word file Start *********************************************

def clean_text(text):
    """Utility function to clean and format text for the document."""
    return text.strip() if text else ''


def generate_questions(request):
    try:
        # Create an in-memory file object
        buffer = BytesIO()

        # Setup document file to save generated word content
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'all_questions_{today}.docx'
        document = Document()

        # Add content to the document
        for question in QuestionBank.objects.all():
            if question.question_sub_type == 'simple_type':
                add_simple_type(question, document)
            elif question.question_sub_type == 'r_and_a_type':
                add_r_and_a_type(question, document)
            elif question.question_sub_type == 'list_type_1':
                add_list_type_1(question, document)
            elif question.question_sub_type == 'list_type_2':
                add_list_type_2(question, document)
            
            # Add a space between questions
            document.add_paragraph()

        # Save the document to the in-memory file object
        document.save(buffer)
        buffer.seek(0)

        # Return the generated file as a downloadable response
        response = FileResponse(buffer, as_attachment=True, filename=file_name)
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)


def add_simple_type(question, document):
    """Add simple type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
    add_options_and_answers(document, question)

def add_r_and_a_type(question, document):
    """Add reason and assertion type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    document.add_paragraph(f"{clean_text(question.question_part_third)}")
    add_options_and_answers(document, question)

def add_list_type_1(question, document):
    """Add list type 1 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    for i in range(1, 9):  # Adjusted to match the model's 8 list rows
        list_row = getattr(question, f'list_1_row{i}', None)
        if list_row:
            document.add_paragraph(f"{i}. {clean_text(list_row)}")
    document.add_paragraph(f"{clean_text(question.question_part_third)}")
    add_options_and_answers(document, question)

def add_list_type_2(question, document):
    """Add list type 2 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    
    # Create table with two columns
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "LIST - I"
    hdr_cells[1].text = "LIST - II"
    
    # Add the list names, if available
    if question.list_1_name:
        hdr_cells[0].text += f"\n({clean_text(question.list_1_name)})"
    if question.list_2_name:
        hdr_cells[1].text += f"\n({clean_text(question.list_2_name)})"

    # Add list rows
    labels = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    for i in range(1, 9):  # Adjusted to match the model's 8 list rows
        list_1_item = clean_text(getattr(question, f'list_1_row{i}', ''))
        list_2_item = clean_text(getattr(question, f'list_2_row{i}', ''))

        if list_1_item or list_2_item:
            row = table.add_row().cells
            row[0].text = f"{labels[i-1]}. {list_1_item}" if list_1_item else ''
            row[1].text = f"{i}. {list_2_item}" if list_2_item else ''

    # Add question details and options below the table
    if question.question_part_third:
        document.add_paragraph(f"{clean_text(question.question_part_third)}")
    
    add_options_and_answers(document, question)

def add_options_and_answers(document, question):
    """Add options and answers to the document."""
    for opt in ['a', 'b', 'c', 'd']:
        option_text = getattr(question, f'answer_option_{opt}', None)
        if option_text:
            document.add_paragraph(f"({opt.lower()}) {clean_text(option_text)}")
    
    document.add_paragraph(f"Correct Answer: {clean_text(question.correct_answer_choice)}")
    document.add_paragraph(f"Solution: {clean_text(question.correct_answer_description)}")
    
    # Format created_at datetime
    created_at_str = question.created_at.strftime('%Y-%m-%d %H:%M:%S')
    document.add_paragraph(f"Created At: {created_at_str}")
    
    # Handle created_by (User object from Django built-in User model)
    if question.created_by:
        created_by_str = question.created_by.get_full_name() or question.created_by.username
        document.add_paragraph(f"Created By: {clean_text(created_by_str)}")



# ************************* Generate Test Word file End *********************************************


# ************************* Generate Clas Plus Word file Start *********************************************

def set_no_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def generate_questions_document(request):
    try:
        # Setup directory and document file to save generated Word file
        base_dir = r'E:/Kamlesh Projects/hajeka_Ias_academy-1.0.0/hajela_ias_academy/media/word_file'
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'class_plus_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in QuestionBank.objects.all():
            # Create main table for the question and set styles
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # Handle question text or assertion and reason
            if question.question_part and question.question_part.strip():
                question_text = question.question_part
            else:
                question_text = (question.question_part_first or '') + "\n" + (question.question_part_third or '')

            # Add question info to the main table
            q_row = table.add_row().cells
            q_row[0].text = 'Question'

            if question.list_1_name and question.list_2_name:
                # Create a sub-table to handle complex question structures
                sub_table = document.add_table(rows=1, cols=2)
                sub_table.style = 'Table Grid'

                # Apply no border style to all cells in the sub-table as they are created
                for cell in sub_table._cells:
                    set_no_border(cell)

                # Modify the headers to include the list names as desired
                sub_hdr_cells = sub_table.rows[0].cells
                sub_hdr_cells[0].text = "LIST - I"
                sub_hdr_cells[1].text = "LIST - II"

                # Add the list names, if available
                if question.list_1_name:
                    sub_hdr_cells[0].text += f"\n({clean_text(question.list_1_name)})"
                if question.list_2_name:
                    sub_hdr_cells[1].text += f"\n({clean_text(question.list_2_name)})"

                # Populate sub-table with list options
                for i in range(1, 9):
                    row_cells = sub_table.add_row().cells
                    for cell in row_cells:
                        set_no_border(cell)  # Ensure borders are removed for new cells too
                    list_1_option = getattr(question, f'list_1_row{i}', '')
                    list_2_option = getattr(question, f'list_2_row{i}', '')
                    row_cells[0].text = f"{chr(64+i)}. {list_1_option}" if list_1_option else ""
                    row_cells[1].text = f"{i}. {list_2_option}" if list_2_option else ""

                # Clear the original cell content, insert question text and sub-table
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run((question.question_part_first or '') + "\n")
                q_row[1]._element.append(sub_table._element)

                # Add 'Codes:' text below the sub-table within the same cell
                p = q_row[1].add_paragraph()
                p.add_run("\nCodes:\t A\t B\t C\t D")

            else:
                # Standard question text handling
                q_row[1].text = question_text

            # Merging cells for question text and image
            q_row[1].merge(q_row[2])
            
            # Handling image insertion if available
            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)
                img_io = BytesIO()
                pil_img.save(img_io, 'JPEG')
                img_io.seek(0)
                paragraph = q_row[1].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(1.5))

            # Table type Questions
            # Check for the presence of table_head_* fields and their corresponding data
            sub_table = None

            table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
            data_fields = [
                [getattr(question, f'head_a_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_b_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_c_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_d_data{j}', None) for j in range(1, 5)]
            ]

            filtered_heads_data = [
                (head, [data for data in datas if data])
                for head, datas in zip(table_heads, data_fields) 
                if getattr(question, head, None) and any(datas)
            ]

            if filtered_heads_data:
                total_rows = 1 + max(len(datas) for _, datas in filtered_heads_data)
                sub_table = document.add_table(rows=total_rows, cols=len(filtered_heads_data))
                sub_table.style = 'Table Grid'
                hdr_cells = sub_table.rows[0].cells
                for idx, (head, _) in enumerate(filtered_heads_data):
                    hdr_cells[idx].text = getattr(question, head, "")
                    paragraph = hdr_cells[idx].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align header text

                for col_idx, (head, data_list) in enumerate(filtered_heads_data):
                    serial_number = 1
                    for row_idx, data in enumerate(data_list):
                        cell = sub_table.cell(row_idx + 1, col_idx)
                        if head == 'table_head_a':
                            cell.text = f"{serial_number}. {data}"
                            serial_number += 1
                        else:
                            cell.text = data

            if sub_table is not None:
                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run((question_text or '') + "\n")
                q_row[1]._element.append(sub_table._element)

            # Additional rows for type, options, solution, and marks
            type_row = table.add_row().cells
            type_row[0].text = 'Type'
            type_row[1].text = question.type_of_question
            type_row[1].merge(type_row[2])

            correct_option_text = ""  # Store text of the correct answer
            valid_options = ['a', 'b', 'c', 'd']  # Include all valid options
            correct_answer = question.correct_answer_choice.lower() if question.correct_answer_choice else None  # Safely handle None

            for opt in valid_options:
                option_text = getattr(question, f"answer_option_{opt}", None)  # Get the option text or None if it doesn't exist
                if option_text:  # Check if the option text is not None
                    opt_row = table.add_row().cells
                    opt_row[0].text = 'Option'
                    opt_row[1].text = f"{opt.upper()}. {option_text}"  # Set the option text
                    is_correct = opt == correct_answer  # Determine if this option is correct
                    opt_row[2].text = 'correct' if is_correct else 'incorrect'  # Set 'correct' or 'incorrect'
                    if is_correct:
                        correct_option_text = option_text  # Store the correct option text

            solution_row = table.add_row().cells
            solution_row[0].text = 'Solution'
            solution_row[1].text = correct_option_text
            solution_row[1].merge(solution_row[2])

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            document.add_paragraph()  # Add space between questions

        document.save(file_path)
        
        # Return the generated file as a downloadable response
        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)



# ************************* Generate Clas Plus Word file Start *********************************************


# ************************* Upload Excel file Start *********************************************

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the uploaded file temporarily
            file = request.FILES['file']
            fs = FileSystemStorage()
            filename = fs.save(file.name, file)
            uploaded_file_url = fs.url(filename)

            # Process the uploaded file
            data = pd.read_excel(fs.path(filename))

            # Replace NaN values with blank strings for text fields and 0 for numeric fields
            data = data.fillna({
                'marks': 0,
                'negative_marks': 0,
                'exam_year': 0,
            }).fillna('')

            # Get the maximum question number from the database
            max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
            if max_question_number:
                start_number = int(max_question_number) + 1
            else:
                start_number = 1

            # Loop through the rows and create QuestionBank entries
            for _, row in data.iterrows():
                while QuestionBank.objects.filter(question_number=str(start_number)).exists():
                    start_number += 1

                try:
                    exam_year = int(row.get('exam_year', 0))
                except ValueError:
                    exam_year = None  # Handle the case where 'exam_year' is not a valid integer

                QuestionBank.objects.create(
                    question_number=str(start_number),
                    type_of_question=row.get('question_type', 'multiple_choice'),
                    exam_name=row.get('exam_name', ''),
                    exam_stage=row.get('exam_stage', ''),
                    exam_year=exam_year,
                    language=row.get('language', ''),
                    script=row.get('script', ''),
                    marks=float(row.get('exam_year1_marks', 0)),
                    negative_marks=float(row.get('exam_year1_negative_marks', 0)),
                    degree_of_difficulty=row.get('degree_of_difficulty', ''),
                    question_sub_type=row.get('question_sub_type', ''),
                    question_part=row.get('question_part', ''),
                    question_part_first=row.get('question_part_first_part', ''),
                    list_1_name=row.get('list_1_name', ''),
                    list_2_name=row.get('list_2_name', ''),
                    list_1_row1=row.get('list_1_row1', ''),
                    list_2_row1=row.get('list_2_row1', ''),
                    list_1_row2=row.get('list_1_row2', ''),
                    list_2_row2=row.get('list_2_row2', ''),
                    list_1_row3=row.get('list_1_row3', ''),
                    list_2_row3=row.get('list_2_row3', ''),
                    list_1_row4=row.get('list_1_row4', ''),
                    list_2_row4=row.get('list_2_row4', ''),
                    list_1_row5=row.get('list_1_row5', ''),
                    list_2_row5=row.get('list_2_row5', ''),
                    list_1_row6=row.get('list_1_row6', ''),
                    list_2_row6=row.get('list_2_row6', ''),
                    list_1_row7=row.get('list_1_row7', ''),
                    list_2_row7=row.get('list_2_row7', ''),
                    list_1_row8=row.get('list_1_row8', ''),
                    list_2_row8=row.get('list_2_row8', ''),
                    question_part_third=row.get('question_part_third_part', ''),
                    answer_option_a=row.get('answer_option_a', ''),
                    answer_option_b=row.get('answer_option_b', ''),
                    answer_option_c=row.get('answer_option_c', ''),
                    answer_option_d=row.get('answer_option_d', ''),
                    correct_answer_choice=row.get('correct_answer_choice', ''),
                    correct_answer_description=row.get('correct_answer_description', ''),
                    subject_name=row.get('subject_name', ''),
                    area_name=row.get('area_name', ''),
                    part_name=row.get('part_name', ''),
                    table_head_a=row.get('table_head_a', ''),
                    table_head_b=row.get('table_head_b', ''),
                    table_head_c=row.get('table_head_c', ''),
                    table_head_d=row.get('table_head_d', ''),
                    head_a_data1=row.get('head_a_data1', ''),
                    head_a_data2=row.get('head_a_data2', ''),
                    head_a_data3=row.get('head_a_data3', ''),
                    head_a_data4=row.get('head_a_data4', ''),
                    head_b_data1=row.get('head_b_data1', ''),
                    head_b_data2=row.get('head_b_data2', ''),
                    head_b_data3=row.get('head_b_data3', ''),
                    head_b_data4=row.get('head_b_data4', ''),
                    head_c_data1=row.get('head_c_data1', ''),
                    head_c_data2=row.get('head_c_data2', ''),
                    head_c_data3=row.get('head_c_data3', ''),
                    head_c_data4=row.get('head_c_data4', ''),
                    head_d_data1=row.get('head_d_data1', ''),
                    head_d_data2=row.get('head_d_data2', ''),
                    head_d_data3=row.get('head_d_data3', ''),
                    head_d_data4=row.get('head_d_data4', '')
                )
                start_number += 1

            messages.success(request, "File uploaded and data processed successfully!")
            return redirect('upload-file')
    else:
        form = UploadFileForm()
    return render(request, 'question_bank/upload.html', {'form': form})

# ************************* Upload Excel file End *********************************************



def get_subjects(request):
    exam_id = request.GET.get('exam_id')
    subjects = Subject.objects.filter(exam_id=exam_id).values('id', 'name')
    return JsonResponse({'subjects': list(subjects)})

def get_areas(request):
    subject_id = request.GET.get('subject_id')
    areas = Area.objects.filter(subject_id=subject_id).values('id', 'name')
    return JsonResponse({'areas': list(areas)})

def get_parts(request):
    area_id = request.GET.get('area_id')
    parts = PartName.objects.filter(area_id=area_id).values('id', 'name')
    return JsonResponse({'parts': list(parts)})


def get_chapters(request):
    part_id = request.GET.get('part_id')
    chapeters = ChapterName.objects.filter(part_id=part_id).values('id', 'name')  # Typo: 'chapeters' should be 'chapters'
    return JsonResponse({'chapters': list(chapeters)})  # Also change to 'chapters'


def get_topics(request):
    chapter_id = request.GET.get('chapter_id')
    topics = TopicName.objects.filter(chapter_id=chapter_id).values('id', 'name')
    return JsonResponse({'topics': list(topics)})



# ************************* Create Simple Type Question Start *********************************************
def add_simple_type_question(request):
    exam_names = ExamName.objects.all()

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''

        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None

        # Initialize the QuestionBank object with the names, not the IDs
        question = QuestionBank(
            question_sub_type=request.POST.get('questionType', 'simple_type'),
            question_part=request.POST.get('question_part_first', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name,  # Saving the name of the exam
            exam_year=exam_year,
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name,  # Saving the name of the subject
            area_name=area_name,  # Saving the name of the area
            part_name=part_name,  # Saving the name of the part
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name=topic_name,  # Saving the name of the topic
            answer_option_a=request.POST.get('answer_option_a', ''),
            answer_option_b=request.POST.get('answer_option_b', ''),
            answer_option_c=request.POST.get('answer_option_c', ''),
            answer_option_d=request.POST.get('answer_option_d', ''),
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        messages.success(request, 'Simple Type Question has been added successfully!')
        return redirect('add-simple-type-question')  # Redirect back to the form

    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/simple_type_form.html', context)

# ************************* Create Simple Type Question End *********************************************



# ************************* Create R and A Type Question Start *********************************************

def add_r_and_a_type_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''
        
        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None


        # Initialize the QuestionBank object
        question = QuestionBank(
            question_sub_type=request.POST.get('questionType', 'r_and_a_type'),
            question_part_first=request.POST.get('question_part_first', ''),
            reason=request.POST.get('reason', ''),
            assertion=request.POST.get('assertion', ''),
            question_part_third= request.POST.get('question_part_third', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name, # Saving the name of the exam
            exam_year=exam_year, # Saving the year for pyqs
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name, # Saving the name of the subject
            area_name=area_name, # Saving the name of the area
            part_name=part_name, # Saving the name of the part
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name=topic_name, # Saving the name of the topic
            answer_option_a=request.POST.get('answer_option_a', ''),
            answer_option_b=request.POST.get('answer_option_b', ''),
            answer_option_c=request.POST.get('answer_option_c', ''),
            answer_option_d=request.POST.get('answer_option_d', ''),
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        messages.success(request, 'R & A Type Question has been added successfully!')
        return redirect('add-r-and-a-type-question')  # Redirect back to the form

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/r_and_a_type_form.html', context)

# ************************* Create R and A Type Question End *********************************************


# ************************* Create List-I Type Question Start *********************************************

def add_list_type_1_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()
 

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''
        
        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None

        # Initialize the QuestionBank object
        question = QuestionBank(
            question_sub_type=request.POST.get('questionType', 'list_type_1'),
            question_part_first=request.POST.get('question_part_first', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name,
            exam_year=exam_year,  # Set to None if empty
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name,
            area_name=area_name,
            part_name=part_name,
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name=topic_name,
            list_1_row1=request.POST.get('list_1_row1', ''),
            list_1_row2=request.POST.get('list_1_row2', ''),
            list_1_row3=request.POST.get('list_1_row3', ''),
            list_1_row4=request.POST.get('list_1_row4', ''),
            list_1_row5=request.POST.get('list_1_row5', ''),
            list_1_row6=request.POST.get('list_1_row6', ''),
            list_1_row7=request.POST.get('list_1_row7', ''),
            list_1_row8=request.POST.get('list_1_row8', ''),
            question_part_third=request.POST.get('question_part_third', ''),
            answer_option_a=request.POST.get('answer_option_a', ''),
            answer_option_b=request.POST.get('answer_option_b', ''),
            answer_option_c=request.POST.get('answer_option_c', ''),
            answer_option_d=request.POST.get('answer_option_d', ''),
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        messages.success(request, 'List-I Type Question has been added successfully!')
        return redirect('add-list-type-1-question')  # Redirect back to the form

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/list_type_1_form.html', context)

# ************************* Create List-I Type Question End *********************************************



# ************************* Create List-II Type Question Start *********************************************
def add_list_type_2_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''
        
        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None


        # Initialize the QuestionBank object
        question = QuestionBank(
            type_of_question='mcq1',
            question_sub_type=request.POST.get('questionType', 'list_type_2'),
            question_part_first=request.POST.get('question_part_first', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name,
            exam_year=exam_year,  # Set to None if empty
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name,
            area_name=area_name,
            part_name=part_name,
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name=topic_name,
            list_1_name=request.POST.get('list_1_name', ''),
            list_2_name=request.POST.get('list_2_name', ''),
            list_1_row1=request.POST.get('list_1_row1', ''),
            list_2_row1=request.POST.get('list_2_row1', ''),
            list_1_row2=request.POST.get('list_1_row2', ''),
            list_2_row2=request.POST.get('list_2_row2', ''),
            list_1_row3=request.POST.get('list_1_row3', ''),
            list_2_row3=request.POST.get('list_2_row3', ''),
            list_1_row4=request.POST.get('list_1_row4', ''),
            list_2_row4=request.POST.get('list_2_row4', ''),
            list_1_row5=request.POST.get('list_1_row5', ''),
            list_2_row5=request.POST.get('list_2_row5', ''),
            question_part_third=request.POST.get('question_part_third', ''),
            answer_option_a=request.POST.get('answer_option_a', ''),
            answer_option_b=request.POST.get('answer_option_b', ''),
            answer_option_c=request.POST.get('answer_option_c', ''),
            answer_option_d=request.POST.get('answer_option_d', ''),
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        messages.success(request, 'List-II Type Question has been added successfully!')
        return redirect('add-list-type-2-question')  # Redirect back to the form

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/list_type_2_form.html', context)

# ************************* Create List-II Type Question End *********************************************


# ************************* Create True and False Type Question Start *********************************************


def add_true_and_false_type_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''
        
        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None

        # Initialize the QuestionBank object with True/False options
        question = QuestionBank(
            question_sub_type=request.POST.get('questionType', 'true_and_false_type'),
            question_part=request.POST.get('question_part_first', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name,
            exam_year=exam_year,
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name,
            area_name=area_name,
            part_name=part_name,
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name = topic_name,
            answer_option_a="True",
            answer_option_b="False",
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        messages.success(request, 'True & False Type Question has been added successfully!')
        return redirect('add-true-and-false-type-question')  # Redirect back to the form

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/true_false_type_form.html', context)

# ************************* Create True and False Type Question End *********************************************


# ************************* Create Fill in the Blank Type Question Start *********************************************

def add_fill_in_the_blank_question(request):
    # Fetch all required data for dropdowns
    exam_names = ExamName.objects.all()

    if request.method == 'POST':
        # Extract form data
        exam_id = request.POST.get('exam_name')
        subject_id = request.POST.get('subject_name')
        area_id = request.POST.get('area_name')
        part_id = request.POST.get('part_name')
        chapter_id = request.POST.get('chapter_name')  # Extract chapter name
        topic_id = request.POST.get('topic_name')

        # Fetch the actual names from the related models
        exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
        subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
        area_name = Area.objects.get(id=area_id).name if area_id else ''
        part_name = PartName.objects.get(id=part_id).name if part_id else ''
        chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

        # Handle topic name (can be selected from dropdown or manually added)
        topic_name = ''
        if topic_id == 'other':
            new_topic_name = request.POST.get('new_topic_name', '')
            if new_topic_name:
                # Ensure the chapter_id is provided when creating a new topic
                topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                topic_name = topic.name
        else:
            topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''
        
        exam_year = request.POST.get('exam_year', None)
        if not exam_year:
            exam_year = None

        # Initialize the QuestionBank object
        question = QuestionBank(
            question_sub_type=request.POST.get('questionType', 'fill_in_the_blank_type'),
            question_part=request.POST.get('question_part_first', ''),
            correct_answer_choice=request.POST.get('correct_answer_choice', ''),
            correct_answer_description=request.POST.get('correct_answer_description', ''),
            exam_name=exam_name,
            exam_year=exam_year,  # Set to None if empty
            marks=float(request.POST.get('marks', 0.0)),
            negative_marks=float(request.POST.get('negative_marks', 0.0)),
            degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
            subject_name=subject_name,
            area_name=area_name,
            part_name=part_name,
            chapter_name=chapter_name,  # Saving the name of the chapter
            topic_name=topic_name,
            answer_option_a=request.POST.get('answer_option_a', ''),
            answer_option_b=request.POST.get('answer_option_b', ''),
            answer_option_c=request.POST.get('answer_option_c', ''),
            answer_option_d=request.POST.get('answer_option_d', ''),
            created_by = request.user  # Assign the logged-in user
        )

        # Save the question to the database
        question.save()

        # Display a success message and redirect back to the form
        messages.success(request, 'Fill in the Blank Question has been added successfully!')
        return redirect('add-fill-in-the-blank-question')

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }

    return render(request, 'question_bank/add_question/fill_in_the_blank_form.html', context)


def add_input_suggestion(request):
    # Fetch all required data for dropdowns
    exam_names = ExamName.objects.all()
    try:
        if request.method == 'POST':
            # Extract form data
            exam_id = request.POST.get('exam_name')
            subject_id = request.POST.get('subject_name')
            area_id = request.POST.get('area_name')
            part_id = request.POST.get('part_name')
            chapter_id = request.POST.get('chapter_name')  # Extract chapter name
            topic_id = request.POST.get('topic_name')

            # Fetch the actual names from the related models
            exam_name = ExamName.objects.get(id=exam_id).name if exam_id else ''
            subject_name = Subject.objects.get(id=subject_id).name if subject_id else ''
            area_name = Area.objects.get(id=area_id).name if area_id else ''
            part_name = PartName.objects.get(id=part_id).name if part_id else ''
            chapter_name = ChapterName.objects.get(id=chapter_id).name if chapter_id else ''  # Get chapter name

            # Handle topic name (can be selected from dropdown or manually added)
            topic_name = ''
            if topic_id == 'other':
                new_topic_name = request.POST.get('new_topic_name', '')
                if new_topic_name:
                    # Ensure the chapter_id is provided when creating a new topic
                    topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter_id=chapter_id)
                    topic_name = topic.name
            else:
                topic_name = TopicName.objects.get(id=topic_id).name if topic_id else ''

            # Extract the main form data
            brief_description = request.POST.get('brief_description')
            details = request.POST.get('details')
            question_video = request.FILES.get('question_video')
            question_link = request.POST.get('question_link')
            other_text = request.POST.get('other_text')

            # Create and save the InputSuggestion object
            suggestion = InputSuggestion(
                brief_description=brief_description,
                details=details,
                question_video=question_video,
                question_link=question_link,
                other_text=other_text,
                exam_name=exam_name,
                subject_name=subject_name,
                area_name=area_name,
                part_name=part_name,
                chapter_name=chapter_name,  # Saving the name of the chapter
                topic_name=topic_name,
                created_by = request.user  # Assign the logged-in user
            )
            suggestion.save()

            # Handle file uploads for images
            if 'question_images' in request.FILES:
                for image in request.FILES.getlist('question_images'):
                    InputSuggestionImage.objects.create(question=suggestion, image=image)

            # Handle file uploads for documents
            if 'question_documents' in request.FILES:
                for document in request.FILES.getlist('question_documents'):
                    InputSuggestionDocument.objects.create(question=suggestion, document=document)

            # Display success message and redirect
            messages.success(request, 'Input Suggestion has been added successfully!')
            return redirect('input-suggestion-list')
        
    except Exception as e:
        # Return the exact error message
        return HttpResponse(f"Error: {str(e)}")
    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
    }
    
    return render(request, 'question_bank/add_input_suggestion.html', context)


def view_input_suggestion(request):
    # Fetch all DescriptiveTypeQuestion entries
    questions = InputSuggestion.objects.all()
    print(questions)

    context = {
        'questions': questions
    }

    return render(request, 'question_bank/input_suggestion_list.html', context)


def question_blog_view(request, question_id):
    # Fetch the question using its ID
    question = get_object_or_404(InputSuggestion, id=question_id)

    context = {
        'question': question
    }
    
    return render(request, 'question_bank/view_input_suggestion.html', context)